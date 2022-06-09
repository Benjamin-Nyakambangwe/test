from docx import Document
from docx.shared import Inches


class Report:
    def __init__(self, company_name, grade_list):
        self.grade_list = grade_list
        self.company_name = company_name

        self.document = Document()

    def generate_report(self):
        self.document.add_heading(
            f'{self.company_name} job description result', 0)

        table = self.document.add_table(rows=1, cols=2)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Position'
        hdr_cells[1].text = 'Grade'

        for job in self.grade_list:
            row_cells = table.add_row().cells
            row_cells[0].text = job[0]
            row_cells[1].text = job[1]

        self.document.add_page_break()

        self.document.add_heading('Grade Justifications', level=1)

        # add table with job descriptions
        justification_table = self.document.add_table(rows=1, cols=4)
        justification_table.style = 'TableGrid'
        hdr_cells2 = justification_table.rows[0].cells
        hdr_cells2[0].text = 'Position'
        hdr_cells2[1].text = 'Grade'
        hdr_cells2[2].text = 'Job Profile'
        hdr_cells2[3].text = 'Justification'

        for job in self.grade_list:
            row_cells = justification_table.add_row().cells
            row_cells[0].text = job[0]
            row_cells[1].text = job[1]
            row_cells[2].text = job[2]
            row_cells[3].text = ""

        self.document.add_page_break()

        return self.document
